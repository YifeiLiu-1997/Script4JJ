from docx import Document
from docx.shared import Pt
import re
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT



def txt2doc(doc_file=None):
    with open('files/你比北京美丽_玖月晞.txt', 'r') as fp:
        for line in fp.readlines():
            line = line[:-1]

            if re.search('☆、chapter [\d]', line):
                paragraph = doc_file.add_paragraph()
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.add_run(line)
                font.name = '新宋体'
                run.bold = True
                run.font.size = Pt(12)  # 设置字体大小
                run.alignment = 1
            else:
                paragraph = doc_file.add_paragraph()
                run = paragraph.add_run(line)
                font = run.font
                # 设置字体样式
                font.name = '新宋体'
                # 设置字体大小
                font.size = Pt(12)

    return doc_file


if __name__ == '__main__':
    doc = Document()
    doc_result = txt2doc(doc)
    doc_result.save('files/你比北京美丽_玖月晞re.docx')
